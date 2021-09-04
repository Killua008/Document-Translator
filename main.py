from kivymd.app import MDApp
from kivy.lang.builder import Builder
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.screenmanager import ScreenManager, Screen
import docx
from googletrans import Translator
import time

Builder_string='''
ScreenManager:
    Main:
    Upload:
    
<Main>:
    name:'main'
    MDLabel:
        text: "Document Translator"
        halign:'center'
        pos_hint:{'center_y':0.9}
        font_style:'H4'
    MDTextField:
        id: input
        hint_text:'Enter the language'
        width:100
        size_hint_x:None
        pos_hint:{'center_y':0.61,'center_x':0.5}

    MDRaisedButton:
        pos_hint:{'center_x':0.5,'center_y':0.5}
        text:'Fetch'
        on_press: app.prints()
    
    MDRectangleFlatButton:
        text: 'Upload'
        pos_hint: {'center_x':0.70,'center_y':0.5}
        on_press: root.manager.current = 'upload'

    MDRectangleFlatButton:
        text: 'Translate'
        pos_hint: {'center_x':0.5,'center_y':0.30}
        on_press: app.translate()
    
    
<Upload>:
    name:'upload'

    FileChooserIconView:
        id:file_select
        on_selection:app.select(*args)
        
    MDRectangleFlatButton:
        text: 'Select'
        pos_hint: {'center_x':0.5,'center_y':0.1}
        on_press: root.manager.current = 'main'

'''

class Main(Screen):
    pass
class Upload(Screen):
    pass

sm=ScreenManager()
sm.add_widget(Main(name='main'))
sm.add_widget(Upload(name='upload'))

class App(MDApp):
    def build(self):
        self.theme_cls.theme_style="Dark"
        self.help_string=Builder.load_string(Builder_string)
        return self.help_string
    def prints(self):
        self.lang=self.help_string.get_screen('main').ids.input.text
        print(self.lang)
    def select(self, *args):
        self.file_path= args[1][0]
        print(self.file_path)

    def translate(self):
        #self.lang=langu
        #self.file_path=file_path
        doc = docx.Document(self.file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        print(len(paragraphs))
        translator = Translator()
        doc = docx.Document()
        for i,para in enumerate(paragraphs):
            try:
                translation = translator.translate(para,src='en',dest=self.lang)
                doc.add_paragraph(translation.text)
                time.sleep(1)
                print("Success "+str(i))
            except:
                print("Error "+str(i))
        doc.save("Translated.docx")
"""    doc=docx.Document(self.file_path)
        fullText= []
        for para in doc.paragraphs:
            fullText.append(para.text)
        sentence=""
        print(sentence)
        translator=Translator()

        with open("texttrans.docx","w",encoding='utf-8')as f:
            out=translator.translate(sentence,dest=self.lang)
            trans=out.text
            f.write(trans)

        out=translator.translate(sentence,dest=self.lang)
        trans=out.text
        print(trans)
"""            

App().run()
